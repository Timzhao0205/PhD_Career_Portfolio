from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from build_abstract_options_no_biomedical import (
    BLUE,
    DARK_BLUE,
    INK,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    PALE_GOLD,
    PALE_GREEN,
    WHITE,
    add_callout,
    add_label_paragraph,
    add_page_field,
    repeat_header,
    set_run_font,
    set_table_geometry,
    setup,
    shade,
    word_count,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "16_Abstract_V1_Analysis_and_Revision.docx"

ORIGINAL_SENTENCES = [
    "Magnetic-field sensors underpin contactless measurements of electrical current, position, motion, material condition, and navigation, yet their selection remains complicated by trade-offs among sensitivity, dynamic range, bandwidth, size, power, cost, and environmental stability.",
    "This review critically compares established and emerging sensor families, including Hall-effect, magnetoresistive, fluxgate, induction, giant magnetoimpedance, superconducting quantum interference, optically pumped, nitrogen-vacancy diamond, magnetoelectric, microelectromechanical, fiber-optic, and resonance-based devices.",
    "Their operating principles, performance envelopes, failure modes, commercial maturity, and representative deployments are assessed across energy, transportation, and industrial manufacturing.",
    "The review also examines data-driven calibration, machine-learning-assisted readout and diagnostics, sensor fusion, and digital twins, together with functional safety, cybersecurity, metrological traceability, and qualification requirements.",
    "The comparison shows that no technology dominates every measurement regime: Hall and magnetoresistive devices lead high-volume deployment, while specialized and quantum approaches offer superior weak-field or spatial performance at greater system complexity.",
    "Future progress will depend increasingly on system integration, trustworthy computation, calibration transfer, and standards-aware design rather than sensitivity alone.",
    "This framework links device physics to application fit and commercialization risk, providing a practical basis for technology selection and research prioritization.",
]

ORIGINAL = " ".join(ORIGINAL_SENTENCES)

ANALYSIS = [
    {
        "role": "Background + central problem",
        "purpose": "Establishes the breadth of magnetic sensing and immediately introduces the selection problem created by competing performance requirements.",
        "works": "It gives the reader a reason for the review and creates useful tension: magnetic sensing is broadly important, but technology selection is not straightforward.",
        "issue": "It carries two long lists in one sentence: five measurement functions and seven selection criteria. The idea is strong, but the sentence is heavier than the opening needs to be.",
        "revision": "Magnetic-field sensors enable contactless measurement of electrical current, position, motion, material condition, and navigation across energy, transportation, and industrial systems.",
        "bridge": "Separate the context from the selection problem. The next sentence can then explain why comparison is necessary.",
        "verdict": "Tighten",
    },
    {
        "role": "Review scope: technologies covered",
        "purpose": "Signals comprehensiveness and identifies the technology taxonomy reviewed in the paper.",
        "works": "It reassures a specialist reader that both established and emerging families are included.",
        "issue": "The twelve-family catalogue is the abstract's largest interruption. It consumes attention without advancing the argument, will date as terminology changes, and duplicates information better carried by the title, keywords, taxonomy figure, and introduction.",
        "revision": "This review compares established and emerging magnetic-sensing technologies by their operating principles, performance limits, failure modes, commercial maturity, and deployment evidence.",
        "bridge": "Merge the useful scope claim with Sentence 3's evaluation method. This turns two inventory sentences into one purposeful method sentence.",
        "verdict": "Replace + merge",
    },
    {
        "role": "Review method + application boundaries",
        "purpose": "Explains the comparison dimensions and limits the application scope to energy, transportation, and industrial manufacturing.",
        "works": "This is essential abstract content because it tells readers how the literature is organized, not merely what technologies are named.",
        "issue": "The passive construction ('are assessed') weakens agency, and 'Their' depends on the long technology list. The sentence is more effective when merged with Sentence 2 and written actively.",
        "revision": "This review compares established and emerging magnetic-sensing technologies by their operating principles, performance limits, failure modes, commercial maturity, and deployment evidence across energy, transportation, and industrial systems.",
        "bridge": "After defining the comparison, move from today's technologies and deployments to the methods and constraints shaping future systems.",
        "verdict": "Keep function; merge",
    },
    {
        "role": "Secondary scope: intelligent methods + deployment constraints",
        "purpose": "Introduces the paper's forward-looking methods and its standards, traceability, and qualification perspective.",
        "works": "This distinguishes the review from a conventional sensor-family catalogue and connects technology to real deployment.",
        "issue": "Two four-item lists compete within one sentence. The relationship between the groups is implicit rather than argumentative.",
        "revision": "It further examines how data-driven calibration, machine-learning-assisted readout, sensor fusion, and digital twins can extend system capability, while safety, cybersecurity, qualification, and metrological traceability constrain adoption.",
        "bridge": "Use 'while' to make the conceptual relationship explicit: intelligent methods create capability, whereas assurance requirements govern adoption.",
        "verdict": "Compress + connect",
    },
    {
        "role": "Principal synthesis / main finding",
        "purpose": "States the review's central comparative result: different technology classes win in different regimes.",
        "works": "This is the strongest evidence-facing sentence because it moves beyond scope and offers a conclusion.",
        "issue": "Naming Hall and magnetoresistive devices is defensible, but it partly recreates the technology-list problem. The abstract needs the market-versus-performance contrast more than the family names.",
        "revision": "The synthesis shows that no technology dominates all measurement regimes: mature integrated sensors best serve many high-volume applications, whereas specialized and quantum approaches offer greater weak-field sensitivity or spatial resolution at higher system complexity.",
        "bridge": "Use 'The synthesis shows' to mark the pivot from what the review covers to what the review concludes.",
        "verdict": "Keep; generalize",
    },
    {
        "role": "Implication + future outlook",
        "purpose": "Interprets the main finding and identifies where future progress is most likely to come from.",
        "works": "It provides a distinctive thesis: system integration and trustworthy computation matter more than sensitivity alone.",
        "issue": "The causal link to Sentence 5 is only implied, and 'calibration transfer' is more specialized than necessary at abstract level.",
        "revision": "Consequently, future progress will depend less on sensitivity alone than on coordinated advances in hardware integration, trustworthy computation, calibration, and standards-aware design.",
        "bridge": "Add 'Consequently' so the outlook is visibly derived from the comparative finding rather than appearing as another topic.",
        "verdict": "Keep; strengthen transition",
    },
    {
        "role": "Contribution + reader value",
        "purpose": "Closes by explaining what the review enables: better technology selection and research prioritization.",
        "works": "It gives both technical and commercialization readers a clear reason to use the review.",
        "issue": "'This framework' has a weak antecedent. The sentence can make the paper's linking function more explicit and active.",
        "revision": "By linking device performance to application requirements and commercialization risk, this review provides a practical framework for technology selection and research prioritization.",
        "bridge": "The opening participial phrase ('By linking...') connects the contribution directly to the analysis performed in the review.",
        "verdict": "Keep; clarify",
    },
]


RECOMMENDED_SENTENCES = [
    "Magnetic-field sensors enable contactless measurement of electrical current, position, motion, material condition, and navigation across energy, transportation, and industrial systems.",
    "Selecting among these technologies remains difficult because sensitivity, dynamic range, bandwidth, environmental stability, size, power, and cost cannot be optimized simultaneously.",
    "This review compares established and emerging magnetic-sensing technologies by their operating principles, performance limits, failure modes, commercial maturity, and deployment evidence.",
    "It further examines how data-driven calibration, machine-learning-assisted readout, sensor fusion, and digital twins can extend system capability, while safety, cybersecurity, qualification, and metrological traceability constrain adoption.",
    "The synthesis shows that no technology dominates all measurement regimes: mature integrated sensors best serve many high-volume applications, whereas specialized and quantum approaches offer greater weak-field sensitivity or spatial resolution at higher system complexity.",
    "Consequently, future progress will depend less on sensitivity alone than on coordinated advances in hardware integration, trustworthy computation, calibration, and standards-aware design.",
    "By linking device performance to application requirements and commercialization risk, this review provides a practical framework for technology selection and research prioritization.",
]

RECOMMENDED = " ".join(RECOMMENDED_SENTENCES)

LEAN = (
    "Magnetic-field sensors support contactless measurement across energy, transportation, and industrial systems, but technology selection remains difficult because sensitivity, range, bandwidth, stability, size, power, and cost cannot be optimized simultaneously. "
    "This review compares established and emerging sensing approaches according to operating principle, performance limits, failure modes, maturity, and deployment evidence. "
    "It also assesses data-driven calibration, machine-learning-assisted readout, sensor fusion, and digital twins, together with the safety, traceability, and qualification requirements that shape adoption. "
    "The evidence shows that no technology dominates all measurement regimes: mature integrated sensors best serve many high-volume applications, whereas specialized and quantum approaches extend weak-field sensitivity or spatial resolution at greater system complexity. "
    "Future progress will therefore depend less on sensitivity alone than on coordinated advances in hardware integration, trustworthy computation, calibration, and standards-aware design. "
    "By connecting device performance with application needs and commercialization risk, this review provides a practical framework for technology selection and research prioritization."
)


def reset_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("ABSTRACT VERSION 1  |  SENTENCE AND FLOW ANALYSIS"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("Non-biomedical scope  |  "), size=8.5, color=MUTED)
    add_page_field(footer)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("ABSTRACT VERSION 1"), size=23, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run("Sentence purpose, flow analysis, and focused revision"), size=14, color=MUTED)
    rows = [
        ("Starting draft", f"178 words; seven sentences; MDPI length compliant"),
        ("Primary diagnosis", "The macro-structure is sound, but the middle is overloaded by consecutive inventories"),
        ("Revision goal", "Preserve the argument while reducing taxonomy detail and making transitions explicit"),
        ("Recommended result", f"{word_count(RECOMMENDED)} words; seven connected sentences"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        shade(table.cell(i, 0), LIGHT_GRAY)
        for run in table.cell(i, 0).paragraphs[0].runs:
            set_run_font(run, color=INK, bold=True)
    set_table_geometry(table, [1900, 7460])


def add_original(doc):
    doc.add_heading("1. Overall summary of Version 1", level=1)
    add_callout(doc, "What the abstract currently says.", "Magnetic sensing is broadly useful but involves difficult performance trade-offs. The review compares technologies and applications, adds intelligent methods and assurance requirements, concludes that no technology wins every regime, and argues that future value will come from integrated, trustworthy, standards-aware systems.", LIGHT_BLUE)
    doc.add_heading("Original abstract", level=2)
    p = doc.add_paragraph(ORIGINAL)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = True
    add_label_paragraph(doc, "Word count:", str(word_count(ORIGINAL)) + " words")
    add_callout(doc, "High-level assessment.", "The seven-sentence architecture is already appropriate: context/problem -> scope/method -> extended scope -> principal synthesis -> implication -> contribution. The weakness is not the structure; it is information density and the lack of explicit connective tissue in the middle.", PALE_GREEN)


def add_argument_map(doc):
    doc.add_heading("2. Purpose of each sentence", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, text in enumerate(("Sentence", "Role in the abstract", "How it should connect", "Decision")):
        table.cell(0, idx).text = text
        shade(table.cell(0, idx), BLUE)
        for run in table.cell(0, idx).paragraphs[0].runs:
            set_run_font(run, size=9.2, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    connections = [
        "Creates the need for comparison.",
        "Answers what the review compares.",
        "Explains how and where comparison is performed.",
        "Extends the review from devices to system adoption.",
        "Pivots from coverage to the main finding.",
        "Derives the future implication from that finding.",
        "Converts the analysis into reader value.",
    ]
    for idx, (entry, connection) in enumerate(zip(ANALYSIS, connections), 1):
        cells = table.add_row().cells
        values = (str(idx), entry["role"], connection, entry["verdict"])
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=8.8)
        if "Replace" in entry["verdict"]:
            shade(cells[3], PALE_GOLD)
        elif "Keep" in entry["verdict"]:
            shade(cells[3], PALE_GREEN)
    set_table_geometry(table, [750, 2500, 3610, 2500])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Context -> selection problem -> review approach -> system layer -> synthesis -> implication -> contribution"), size=10, color=DARK_BLUE, bold=True)


def add_sentence_analysis(doc):
    doc.add_heading("3. Sentence-by-sentence analysis", level=1)
    for idx, (sentence, entry) in enumerate(zip(ORIGINAL_SENTENCES, ANALYSIS), 1):
        doc.add_heading(f"Sentence {idx}: {entry['role']}", level=2)
        add_callout(doc, "Original.", sentence, LIGHT_GRAY)
        add_label_paragraph(doc, "Purpose:", entry["purpose"], color=INK)
        add_label_paragraph(doc, "What works:", entry["works"], color=DARK_BLUE)
        add_label_paragraph(doc, "Flow or detail issue:", entry["issue"], color="7A5A00")
        add_label_paragraph(doc, "Connection strategy:", entry["bridge"], color=DARK_BLUE)
        add_callout(doc, "Proposed sentence.", entry["revision"], PALE_GREEN)


def add_flow_diagnosis(doc):
    doc.add_page_break()
    doc.add_heading("4. How to improve the flow", level=1)
    doc.add_heading("Current flow problem", level=2)
    add_label_paragraph(doc, "Inventory accumulation:", "Sentences 2-4 successively list sensor families, comparison dimensions, future methods, and assurance topics. The reader processes coverage rather than a developing argument.")
    add_label_paragraph(doc, "Weak referents:", "'Their' in Sentence 3 points back to a long catalogue, while 'This framework' in Sentence 7 has no single clear antecedent.")
    add_label_paragraph(doc, "Under-signaled pivot:", "Sentence 5 contains the main finding, but the transition from scope to synthesis can be made clearer.")
    doc.add_heading("Revised connective logic", level=2)
    strategies = [
        ("Split context from tension", "Let the first sentence establish importance and the second explain why selection is difficult."),
        ("Merge scope with method", "Replace the technology catalogue with a general scope phrase and immediately state the comparison dimensions."),
        ("Express relationships, not lists", "Use 'while' to contrast capability-enhancing methods with adoption-constraining assurance requirements."),
        ("Signal the result", "Use 'The synthesis shows' to tell the reader that the abstract has moved from method to conclusion."),
        ("Signal causality", "Use 'Consequently' to make the future outlook a direct implication of the comparative finding."),
        ("Close actively", "Use 'By linking...' to specify how the review creates a practical framework."),
    ]
    table = doc.add_table(rows=1, cols=2)
    for idx, text in enumerate(("Technique", "Effect on flow")):
        table.cell(0, idx).text = text
        shade(table.cell(0, idx), BLUE)
        for run in table.cell(0, idx).paragraphs[0].runs:
            set_run_font(run, size=9.5, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    for name, effect in strategies:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = effect
        shade(cells[0], LIGHT_GRAY)
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=9.2)
    set_table_geometry(table, [2600, 6760])


def add_detail_reduction(doc):
    doc.add_heading("5. What detail belongs outside the abstract", level=1)
    add_callout(doc, "Main recommendation.", "Do not enumerate every magnetic-sensing family in the abstract. Comprehensiveness should be demonstrated by the paper's taxonomy, comparison table, title, keywords, and introduction, not by a twelve-item sentence.", PALE_GOLD)
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(("Current detail", "Abstract-level replacement", "Best location for full detail")):
        table.cell(0, idx).text = text
        shade(table.cell(0, idx), BLUE)
        for run in table.cell(0, idx).paragraphs[0].runs:
            set_run_font(run, size=9.2, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    rows = [
        ("Twelve named sensor families", "Established and emerging magnetic-sensing technologies", "Introduction taxonomy paragraph; Section 2 headings; comparison table"),
        ("Five operating attributes plus deployments", "Performance limits, maturity, and deployment evidence", "Methods/scope paragraph and master comparison table"),
        ("Four future methods plus four assurance topics", "Intelligent sensing methods and adoption requirements", "Future-methods and standards sections"),
        ("Hall and magnetoresistive market leadership", "Mature integrated sensors in high-volume applications", "Technology and application synthesis sections"),
        ("Calibration transfer", "Calibration", "Detailed future-methods discussion"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=8.8)
    set_table_geometry(table, [3000, 3000, 3360])


def add_revisions(doc):
    doc.add_page_break()
    doc.add_heading("6. Recommended revision", level=1)
    add_callout(doc, "Recommendation.", "Use the 166-word version as the next working draft. It preserves the original seven-sentence logic, removes the technology catalogue, and makes every transition explicit.", PALE_GREEN)
    p = doc.add_paragraph(RECOMMENDED)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = True
    add_label_paragraph(doc, "Word count:", f"{word_count(RECOMMENDED)} words")
    add_label_paragraph(doc, "Keywords:", "magnetic-field sensors; sensor performance; energy systems; transportation; industrial sensing; machine learning; digital twins")
    doc.add_heading("Sentence functions in the revised version", level=2)
    revised_roles = [
        "Context and application relevance",
        "Selection problem and motivation",
        "Review scope and comparison method",
        "Future-system enablers versus adoption constraints",
        "Principal comparative finding",
        "Implication and future direction",
        "Contribution and practical value",
    ]
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(("Sentence", "Function", "Transition cue")):
        table.cell(0, idx).text = text
        shade(table.cell(0, idx), BLUE)
        for run in table.cell(0, idx).paragraphs[0].runs:
            set_run_font(run, size=9.2, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    cues = ["Topic established", "Selecting...", "This review...", "It further... while...", "The synthesis shows... whereas...", "Consequently...", "By linking..."]
    for idx, (role, cue) in enumerate(zip(revised_roles, cues), 1):
        cells = table.add_row().cells
        for cell, value in zip(cells, (str(idx), role, cue)):
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=9.0)
    set_table_geometry(table, [900, 5160, 3300])

    doc.add_heading("Optional leaner version", level=2)
    add_callout(doc, "Use if you want less catalogue-like detail throughout.", "This version removes the opening list of measurement functions as well as the technology-family list. It is more fluid but slightly less concrete.", LIGHT_BLUE)
    p = doc.add_paragraph(LEAN)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = True
    add_label_paragraph(doc, "Word count:", f"{word_count(LEAN)} words")


def validate():
    assert word_count(ORIGINAL) == 178
    assert word_count(RECOMMENDED) <= 200
    assert word_count(LEAN) <= 200
    assert len(ORIGINAL_SENTENCES) == len(ANALYSIS) == len(RECOMMENDED_SENTENCES) == 7
    for abstract in (RECOMMENDED, LEAN):
        lower = abstract.lower()
        for term in ("biomedical", "clinical", "biosensor", "magnetoencephal", "magnetocardi"):
            assert term not in lower


def build():
    validate()
    doc = Document()
    setup(doc)
    reset_header_footer(doc)
    add_masthead(doc)
    add_original(doc)
    add_argument_map(doc)
    add_sentence_analysis(doc)
    add_flow_diagnosis(doc)
    add_detail_reduction(doc)
    add_revisions(doc)
    props = doc.core_properties
    props.title = "Abstract Version 1: Sentence and Flow Analysis"
    props.subject = "Purpose map, flow diagnosis, detail reduction, and revised abstracts"
    props.author = "Research drafting brief"
    props.keywords = "magnetic-field sensors; abstract revision; sentence purpose; flow"
    doc.save(OUT)
    print(OUT)
    print(f"original={word_count(ORIGINAL)} recommended={word_count(RECOMMENDED)} lean={word_count(LEAN)}")


if __name__ == "__main__":
    build()
