from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "15_Abstract_Options_No_Biomedical.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5D6875"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
PALE_GREEN = "E7F2EA"
PALE_GOLD = "FFF4D6"
WHITE = "FFFFFF"


ABSTRACTS = [
    {
        "name": "Version 1 - Balanced and comprehensive (recommended)",
        "positioning": "Best match to the current full review plan; balances technologies, three application domains, intelligent methods, and standards.",
        "title": "Magnetic-Field Sensing Technologies for Energy, Transportation, and Industry: Principles, Applications, Intelligence, and Standards",
        "abstract": "Magnetic-field sensors underpin contactless measurements of electrical current, position, motion, material condition, and navigation, yet their selection remains complicated by trade-offs among sensitivity, dynamic range, bandwidth, size, power, cost, and environmental stability. This review critically compares established and emerging sensor families, including Hall-effect, magnetoresistive, fluxgate, induction, giant magnetoimpedance, superconducting quantum interference, optically pumped, nitrogen-vacancy diamond, magnetoelectric, microelectromechanical, fiber-optic, and resonance-based devices. Their operating principles, performance envelopes, failure modes, commercial maturity, and representative deployments are assessed across energy, transportation, and industrial manufacturing. The review also examines data-driven calibration, machine-learning-assisted readout and diagnostics, sensor fusion, and digital twins, together with functional safety, cybersecurity, metrological traceability, and qualification requirements. The comparison shows that no technology dominates every measurement regime: Hall and magnetoresistive devices lead high-volume deployment, while specialized and quantum approaches offer superior weak-field or spatial performance at greater system complexity. Future progress will depend increasingly on system integration, trustworthy computation, calibration transfer, and standards-aware design rather than sensitivity alone. This framework links device physics to application fit and commercialization risk, providing a practical basis for technology selection and research prioritization.",
        "keywords": "magnetic-field sensors; Hall effect; magnetoresistance; sensor fusion; machine learning; digital twins; functional safety",
    },
    {
        "name": "Version 2 - Technology and performance centered",
        "positioning": "Best for a conventional Sensors review whose main contribution is the cross-family comparison and technology-selection framework.",
        "title": "From Hall Sensors to Quantum Magnetometry: A Comparative Review for Energy, Transportation, and Industrial Systems",
        "abstract": "Magnetic-field sensing spans measurement regimes from high-current power conversion and automotive position detection to precision navigation, condition monitoring, and weak-field metrology. Because sensitivity alone does not determine application fitness, this review compares magnetic-sensor technologies through a common set of figures of merit: field range, noise, bandwidth, linearity, offset and drift, temperature tolerance, power, size, cost, and maturity. Hall-effect, anisotropic, giant, and tunneling magnetoresistive sensors are examined alongside fluxgate, induction, giant magnetoimpedance, superconducting quantum interference, optically pumped, nitrogen-vacancy diamond, magnetoelectric, microelectromechanical, fiber-optic, and resonance-based approaches. The literature is synthesized by operating principle, achievable performance, dominant failure modes, manufacturability, and deployment evidence in energy, transportation, and industrial manufacturing. Cross-family comparison confirms that broad dynamic range, integration, and qualification sustain Hall and magnetoresistive technologies in volume markets, whereas specialized and quantum platforms occupy narrower regimes requiring higher system complexity. Emerging advances in learned calibration, array processing, sensor fusion, and digital twins can extend usable performance but introduce new questions of robustness and traceability. The resulting capability map clarifies where physical limits, implementation choices, and commercial readiness constrain technology substitution and identifies research priorities for reliable next-generation magnetic sensing.",
        "keywords": "magnetic sensors; Hall-effect sensors; magnetoresistive sensors; quantum magnetometry; performance comparison; technology readiness",
    },
    {
        "name": "Version 3 - Applications and deployment centered",
        "positioning": "Best if the paper is marketed primarily to readers in electrification, mobility, automation, and asset monitoring.",
        "title": "Magnetic-Field Sensors in Energy, Transportation, and Industrial Manufacturing: Technologies, Deployment, and Outlook",
        "abstract": "Electrification, automated mobility, and intelligent manufacturing increasingly depend on magnetic sensing for isolated current measurement, position and speed detection, navigation, non-destructive testing, and equipment health monitoring. This review evaluates how application requirements determine the choice among Hall-effect, magnetoresistive, fluxgate, induction, giant magnetoimpedance, superconducting quantum interference, optically pumped, nitrogen-vacancy diamond, magnetoelectric, microelectromechanical, fiber-optic, and resonance-based sensors. Evidence is organized across three domains: energy systems, including power converters, batteries, grids, and renewable assets; transportation, including automotive actuation, traction systems, and navigation; and industrial manufacturing, including inspection, condition monitoring, robotics, and process metrology. For each domain, the review links field range, resolution, bandwidth, temperature, power, packaging, safety, and cost constraints to the technologies that are established, emerging, or unsuitable. Hall and magnetoresistive devices dominate many deployed systems because they combine integration, robustness, and qualification, while specialized technologies provide advantages where weak-field sensitivity, isolation, or spatial resolution justifies greater complexity. Data-driven calibration, predictive diagnostics, sensor fusion, and digital twins are assessed as cross-domain enablers. The review concludes that deployment success depends on the complete sensing chain, including calibration, algorithms, standards, and lifecycle assurance, rather than on the sensing element alone.",
        "keywords": "magnetic-field sensing; energy systems; transportation sensors; industrial monitoring; current sensing; predictive maintenance; sensor fusion",
    },
    {
        "name": "Version 4 - Intelligent sensing and digital-twin centered",
        "positioning": "Best for an artificial-intelligence or smart-sensing special issue, while retaining a rigorous hardware foundation.",
        "title": "Magnetic-Field Sensors from Devices to Intelligent Systems: Machine Learning, Fusion, Digital Twins, and Industrial Applications",
        "abstract": "Magnetic-sensor performance is increasingly determined not only by the transducer, but also by calibration, signal processing, fusion, diagnostics, and lifecycle models. This review connects the physics and maturity of established and emerging magnetic-field sensors with the computational methods that are reshaping their use in energy, transportation, and industrial manufacturing. Hall-effect, magnetoresistive, fluxgate, induction, giant magnetoimpedance, superconducting quantum interference, optically pumped, nitrogen-vacancy diamond, magnetoelectric, microelectromechanical, fiber-optic, and resonance-based devices are compared in terms of sensitivity, range, bandwidth, drift, environmental tolerance, integration, cost, and readiness. The review then examines data-driven calibration and compensation, machine-learning-assisted readout and fault diagnosis, magnetic-inertial fusion, sensor arrays, source reconstruction, and digital twins. Available evidence indicates that learning can improve calibration transfer, denoising, and condition classification, but does not remove front-end noise, offset, or hardware constraints. Machine-level digital twins are more mature than sensor-level twins that explicitly track calibration state and uncertainty. Trustworthy adoption therefore requires representative data, physics-aware models, distribution-shift testing, metrological traceability, cybersecurity, and standards-compatible change control. By separating demonstrated capabilities from outlook claims, this review defines a practical research agenda for intelligent magnetic sensing without overstating algorithmic readiness.",
        "keywords": "magnetic sensors; machine learning; data-driven calibration; sensor fusion; digital twins; condition monitoring; trustworthy sensing",
    },
    {
        "name": "Version 5 - Standards and commercialization centered",
        "positioning": "Best when standards, qualification, technology readiness, and investment risk are intended as the paper's differentiator.",
        "title": "Magnetic-Field Sensors from Laboratory Performance to Qualified Systems: Technologies, Standards, and Commercial Readiness",
        "abstract": "Magnetic-field sensor research produces devices with widely differing sensitivity, bandwidth, operating conditions, manufacturability, and commercial maturity, making laboratory performance an incomplete predictor of deployment success. This review compares established and emerging magnetic-sensor families and evaluates their transition into energy, transportation, and industrial manufacturing systems. Hall-effect, magnetoresistive, fluxgate, induction, giant magnetoimpedance, superconducting quantum interference, optically pumped, nitrogen-vacancy diamond, magnetoelectric, microelectromechanical, fiber-optic, and resonance-based technologies are assessed through operating principle, performance envelope, failure modes, supply-chain depth, and representative use. The analysis distinguishes mass-market platforms from established niches and pioneering technologies, while treating market estimates and technology readiness as attributed evidence or explicit author assessment. Functional safety, electromagnetic compatibility, cybersecurity, industrial-control requirements, calibration, and metrological traceability are examined as market gates rather than administrative afterthoughts. Data-driven calibration, machine learning, fusion, and digital twins may reduce calibration burden and create diagnostic value, but they also complicate assurance, uncertainty, and change control. The review finds that successful commercialization depends on reproducible system performance, qualification strategy, integration cost, and lifecycle support as much as on sensor sensitivity. It provides a structured basis for technical due diligence, product planning, and standards-aware research investment.",
        "keywords": "magnetic-field sensors; technology readiness; functional safety; qualification; metrological traceability; commercialization; industrial sensing",
    },
    {
        "name": "Version 6 - Concise and broad-audience",
        "positioning": "Best as a clean starting point when the final manuscript emphasis is still evolving; shortest and easiest to refine later.",
        "title": "Magnetic-Field Sensors for Energy, Transportation, and Industry: A Review of Technologies and Emerging Intelligent Systems",
        "abstract": "Magnetic-field sensors enable contactless measurement of electrical current, position, motion, material condition, and navigation across modern energy, transportation, and industrial systems. This review compares established and emerging sensor families using common criteria that include sensitivity, range, bandwidth, drift, temperature tolerance, size, power, cost, and commercial maturity. It relates these performance trade-offs to current sensing in power electronics and batteries, automotive and transportation control, navigation, non-destructive testing, condition monitoring, robotics, and process metrology. The review further assesses data-driven calibration, machine-learning-assisted readout and diagnostics, sensor fusion, and digital twins, together with the standards and traceability requirements that govern critical deployments. No technology is optimal for every measurement regime: Hall-effect and magnetoresistive devices dominate many high-volume applications, while specialized and quantum sensors provide advantages where exceptional weak-field sensitivity, isolation, or spatial resolution justifies added complexity. The evidence indicates that future gains will increasingly arise from integrated sensing systems that combine appropriate hardware with trustworthy algorithms, calibration, diagnostics, and qualification. This perspective provides a practical framework for selecting technologies, identifying research gaps, and evaluating commercial readiness.",
        "keywords": "magnetic-field sensors; energy; transportation; industrial sensing; machine learning; digital twins",
    },
]


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w]+(?:[-'][\w]+)*\b", text, flags=re.UNICODE))


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("ABSTRACT DEVELOPMENT BRIEF  |  MDPI SENSORS"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("Non-biomedical scope  |  "), size=8.5, color=MUTED)
    add_page_field(footer)


def add_label_paragraph(doc, label, text, after=6, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = True
    r = p.add_run(label + " ")
    set_run_font(r, color=color, bold=True)
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(label + " ")
    set_run_font(r, color=INK, bold=True)
    p.add_run(text)
    return table


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("ABSTRACT OPTIONS")
    set_run_font(r, size=23, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Magnetic-field sensors review - revised three-domain scope")
    set_run_font(r, size=14, color=MUTED)
    rows = [
        ("Target journal", "Sensors (MDPI) - Review"),
        ("Scope", "Energy, transportation, and industrial/manufacturing applications; biomedical applications excluded"),
        ("Requirements checked", "27 July 2026 against Sensors Instructions for Authors and the MDPI Layout Style Guide"),
        ("Recommendation", "Begin with Version 1; use Version 4 or 5 only if the manuscript's differentiating emphasis changes"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        shade(table.cell(i, 0), LIGHT_GRAY)
        for run in table.cell(i, 0).paragraphs[0].runs:
            set_run_font(run, color=INK, bold=True)
    set_table_geometry(table, [1800, 7560])


def add_requirements(doc):
    doc.add_heading("1. Submission constraints applied", level=1)
    add_callout(doc, "Safe drafting rule.", "Each option is a single self-contained paragraph of 200 words or fewer. The internal flow is background and purpose, review approach, principal synthesis, and conclusion, without visible subheadings.")
    table = doc.add_table(rows=1, cols=3)
    for idx, value in enumerate(("Requirement", "Applied here", "Drafting implication")):
        table.cell(0, idx).text = value
        shade(table.cell(0, idx), BLUE)
        for run in table.cell(0, idx).paragraphs[0].runs:
            set_run_font(run, size=9.5, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    rows = [
        ("Length", "200 words maximum", "All versions have a printed, programmatically checked word count."),
        ("Form", "One paragraph; structured logic without headings", "No Background/Methods/Results/Conclusion labels appear inside an abstract."),
        ("Independence", "Self-contained", "No citations, figures, tables, links, undefined abbreviations, or claims absent from the planned review."),
        ("Keywords", "3-10 pertinent terms", "Each option uses 6-7 searchable keywords immediately after the abstract."),
        ("Review fit", "Concise update on progress and gaps", "Each option states the comparison method, main synthesis, and implications."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=9.2)
    set_table_geometry(table, [1700, 2200, 5460])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Official guidance checked: ")
    set_run_font(p.runs[0], size=9.5, color=MUTED, bold=True)
    r = p.add_run("https://www.mdpi.com/journal/Sensors/instructions and https://www.mdpi.com/authors/layout")
    set_run_font(r, size=9.5, color=MUTED)


def add_scope_note(doc):
    doc.add_heading("2. Scope decision: biomedical applications excluded", level=1)
    add_label_paragraph(doc, "Included application domains:", "energy; transportation; industrial and manufacturing systems.")
    add_label_paragraph(doc, "Excluded from the abstracts:", "magnetoencephalography, magnetocardiography, magnetic particle imaging, biosensors, medical tracking, medical-device standards, and clinical claims.")
    add_label_paragraph(doc, "Retained technology families:", "superconducting quantum interference devices, optically pumped magnetometers, nitrogen-vacancy diamond sensors, and other specialized platforms remain in the technology taxonomy, but are framed through weak-field metrology, navigation, inspection, power, and industrial uses rather than biomedical use.")
    add_callout(doc, "Manuscript consistency action.", "The final title, outline, figures, application matrix, standards section, and conclusion should later be revised to remove biomedical promises. The abstract options below already implement that scope change.", PALE_GOLD)


def add_abstracts(doc):
    doc.add_heading("3. Abstract options", level=1)
    for index, item in enumerate(ABSTRACTS, 1):
        if index > 1:
            doc.add_page_break()
        doc.add_heading(item["name"], level=2)
        add_label_paragraph(doc, "Use when:", item["positioning"], color=DARK_BLUE)
        add_label_paragraph(doc, "Candidate title:", item["title"], color=INK)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.10
        p.paragraph_format.keep_together = True
        r = p.add_run("Abstract")
        set_run_font(r, size=11, color=BLUE, bold=True)
        p = doc.add_paragraph(item["abstract"])
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.10
        p.paragraph_format.keep_together = True
        add_label_paragraph(doc, "Keywords:", item["keywords"], color=DARK_BLUE)
        count = word_count(item["abstract"])
        status = "PASS" if count <= 200 else "FAIL"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"MDPI length check: {count} words - {status}")
        set_run_font(r, size=9.5, color=("1F3A5F" if status == "PASS" else "9B1C1C"), bold=True)


def add_comparison(doc):
    doc.add_page_break()
    doc.add_heading("4. Selection guide", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, value in enumerate(("Version", "Primary strength", "Main tradeoff", "Recommendation")):
        table.cell(0, idx).text = value
        shade(table.cell(0, idx), BLUE)
        for run in table.cell(0, idx).paragraphs[0].runs:
            set_run_font(run, size=9.2, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    rows = [
        ("1", "Most complete representation of the planned review", "Dense technology list", "Recommended base"),
        ("2", "Strongest technology-comparison identity", "Less explicit business/standards emphasis", "Use for conventional review framing"),
        ("3", "Strongest application relevance", "Technology novelty is less prominent", "Use for applied readership"),
        ("4", "Strongest intelligent-sensing hook", "Could over-shape expectations toward AI", "Use for AI/smart-sensing venue"),
        ("5", "Strongest standards and commercialization distinction", "Less conventional academic tone", "Use if business angle is central"),
        ("6", "Cleanest and shortest", "Fewer specific family cues", "Best fallback or cover-letter synopsis"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=8.8)
        if row[0] == "1":
            shade(cells[3], PALE_GREEN)
    set_table_geometry(table, [800, 2960, 2800, 2800])
    doc.add_heading("Recommended next edit", level=2)
    add_callout(doc, "Start from Version 1.", "After the detailed outline is revised to three application domains, shorten the device-family list if necessary and replace the broad synthesis statements with the manuscript's final, fully supported conclusions. Keep the abstract below 200 words during every revision.", PALE_GREEN)


def validate():
    for item in ABSTRACTS:
        count = word_count(item["abstract"])
        if count > 200:
            raise ValueError(f"{item['name']} exceeds 200 words: {count}")
        if "\n" in item["abstract"]:
            raise ValueError(f"{item['name']} is not a single paragraph")
        lower = item["abstract"].lower()
        prohibited = ("biomedical", "medical", "clinical", "magnetoencephal", "magnetocardi", "biosensor")
        if any(term in lower for term in prohibited):
            raise ValueError(f"{item['name']} contains excluded-scope language")
        keyword_count = len([x for x in item["keywords"].split(";") if x.strip()])
        if not 3 <= keyword_count <= 10:
            raise ValueError(f"{item['name']} keyword count is {keyword_count}")


def build():
    validate()
    doc = Document()
    setup(doc)
    add_masthead(doc)
    add_requirements(doc)
    add_scope_note(doc)
    add_abstracts(doc)
    add_comparison(doc)
    props = doc.core_properties
    props.title = "Abstract Options for Magnetic-Field Sensors Review - Non-Biomedical Scope"
    props.subject = "Six MDPI Sensors-compliant abstract drafts"
    props.author = "Research drafting brief"
    props.keywords = "magnetic-field sensors; abstract; MDPI Sensors; review"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
    for item in ABSTRACTS:
        print(item["name"], word_count(item["abstract"]))


if __name__ == "__main__":
    build()
