from __future__ import annotations

import sys
import zipfile
from lxml import etree
from docx import Document

from build_abstract_v1_analysis import (
    ANALYSIS,
    LEAN,
    ORIGINAL,
    ORIGINAL_SENTENCES,
    OUT,
    RECOMMENDED,
    RECOMMENDED_SENTENCES,
    word_count,
)


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
errors = []

with zipfile.ZipFile(OUT) as archive:
    bad = archive.testzip()
    if bad:
        errors.append(f"Corrupt DOCX member: {bad}")
    document_xml = etree.fromstring(archive.read("word/document.xml"))
    styles_xml = etree.fromstring(archive.read("word/styles.xml"))

doc = Document(OUT)
paragraphs = [p.text for p in doc.paragraphs]
all_text = "\n".join(paragraphs)
for table in doc.tables:
    for row in table.rows:
        all_text += "\n" + " | ".join(cell.text for cell in row.cells)

for token in ("turn0", "turn1", "codex-file-citation", "{{", "}}", "TODO", "TBD"):
    if token in all_text:
        errors.append(f"Internal or placeholder token present: {token}")

for sentence in ORIGINAL_SENTENCES:
    if sentence not in all_text:
        errors.append("Missing original sentence: " + sentence[:60])
for sentence in RECOMMENDED_SENTENCES:
    if sentence not in RECOMMENDED:
        errors.append("Recommended sentence not assembled: " + sentence[:60])
for label, abstract in (("original", ORIGINAL), ("recommended", RECOMMENDED), ("lean", LEAN)):
    count = word_count(abstract)
    if label != "original" and count > 200:
        errors.append(f"{label} abstract exceeds 200 words: {count}")
    if label != "original":
        lower = abstract.lower()
        for term in ("biomedical", "clinical", "biosensor", "magnetoencephal", "magnetocardi"):
            if term in lower:
                errors.append(f"Excluded-scope term '{term}' in {label} abstract")

sects = document_xml.xpath("//w:sectPr", namespaces=NS)
for idx, sect in enumerate(sects, 1):
    pg_sz = sect.find("w:pgSz", NS)
    pg_mar = sect.find("w:pgMar", NS)
    if pg_sz is None or pg_sz.get(f"{{{NS['w']}}}w") != "12240" or pg_sz.get(f"{{{NS['w']}}}h") != "15840":
        errors.append(f"Section {idx}: page size mismatch")
    if pg_mar is None:
        errors.append(f"Section {idx}: page margins missing")
    else:
        for side in ("top", "right", "bottom", "left"):
            if pg_mar.get(f"{{{NS['w']}}}{side}") != "1440":
                errors.append(f"Section {idx}: {side} margin mismatch")

styles = {s.get(f"{{{NS['w']}}}styleId"): s for s in styles_xml.xpath("//w:style", namespaces=NS)}
for style_id in ("Normal", "Heading1", "Heading2", "Heading3"):
    if style_id not in styles:
        errors.append(f"Missing required style: {style_id}")

tables = document_xml.xpath("//w:tbl", namespaces=NS)
for idx, table in enumerate(tables, 1):
    tbl_w = table.find("w:tblPr/w:tblW", NS)
    tbl_ind = table.find("w:tblPr/w:tblInd", NS)
    widths = [int(x.get(f"{{{NS['w']}}}w")) for x in table.findall("w:tblGrid/w:gridCol", NS)]
    if tbl_w is None or int(tbl_w.get(f"{{{NS['w']}}}w")) != sum(widths):
        errors.append(f"Table {idx}: tblW/grid mismatch")
    if tbl_ind is None or tbl_ind.get(f"{{{NS['w']}}}w") != "120":
        errors.append(f"Table {idx}: tblInd mismatch")
    for ridx, row in enumerate(table.findall("w:tr", NS), 1):
        cells = row.findall("w:tc", NS)
        cell_widths = []
        for cell in cells:
            tcw = cell.find("w:tcPr/w:tcW", NS)
            cell_widths.append(int(tcw.get(f"{{{NS['w']}}}w")) if tcw is not None else -1)
        if cell_widths != widths:
            errors.append(f"Table {idx} row {ridx}: tcW/grid mismatch")
        height = row.find("w:trPr/w:trHeight", NS)
        if height is not None and height.get(f"{{{NS['w']}}}hRule") == "exact":
            errors.append(f"Table {idx} row {ridx}: fixed exact height")

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sentence_analyses={len(ANALYSIS)}")
print(f"original={word_count(ORIGINAL)} recommended={word_count(RECOMMENDED)} lean={word_count(LEAN)}")
if errors:
    print("AUDIT FAILED")
    for error in errors:
        print("-", error)
    sys.exit(1)
print("AUDIT PASSED")
