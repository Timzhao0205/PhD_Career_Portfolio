from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from build_abstract_options_no_biomedical import ABSTRACTS, OUT, word_count


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
errors = []

with zipfile.ZipFile(OUT) as archive:
    bad = archive.testzip()
    if bad:
        errors.append(f"Corrupt DOCX member: {bad}")
    document_xml = etree.fromstring(archive.read("word/document.xml"))
    styles_xml = etree.fromstring(archive.read("word/styles.xml"))

doc = Document(OUT)
paragraph_text = [p.text for p in doc.paragraphs]
all_text = "\n".join(paragraph_text)

for token in ("turn0", "turn1", "codex-file-citation", "{{", "}}", "TODO", "TBD"):
    if token in all_text:
        errors.append(f"Internal or placeholder token present: {token}")

for item in ABSTRACTS:
    abstract = item["abstract"]
    count = word_count(abstract)
    if count > 200:
        errors.append(f"Over-length abstract: {item['name']} ({count})")
    if paragraph_text.count(abstract) != 1:
        errors.append(f"Abstract not found exactly once as one paragraph: {item['name']}")
    lower = abstract.lower()
    for term in ("biomedical", "medical", "clinical", "magnetoencephal", "magnetocardi", "biosensor"):
        if term in lower:
            errors.append(f"Excluded scope term '{term}' in {item['name']}")
    keyword_count = len([x for x in item["keywords"].split(";") if x.strip()])
    if not 3 <= keyword_count <= 10:
        errors.append(f"Keyword count outside 3-10: {item['name']} ({keyword_count})")

sects = document_xml.xpath("//w:sectPr", namespaces=NS)
for idx, sect in enumerate(sects, 1):
    pg_sz = sect.find("w:pgSz", NS)
    pg_mar = sect.find("w:pgMar", NS)
    if pg_sz is None or pg_sz.get(f"{{{NS['w']}}}w") != "12240" or pg_sz.get(f"{{{NS['w']}}}h") != "15840":
        errors.append(f"Section {idx}: page size mismatch")
    if pg_mar is None:
        errors.append(f"Section {idx}: page margins missing")
    else:
        for side in ("top", "right", "bottom", "left"):
            if pg_mar.get(f"{{{NS['w']}}}{side}") != "1440":
                errors.append(f"Section {idx}: {side} margin mismatch")

styles = {s.get(f"{{{NS['w']}}}styleId"): s for s in styles_xml.xpath("//w:style", namespaces=NS)}
for style_id in ("Normal", "Heading1", "Heading2", "Heading3"):
    if style_id not in styles:
        errors.append(f"Missing required style: {style_id}")

tables = document_xml.xpath("//w:tbl", namespaces=NS)
for idx, table in enumerate(tables, 1):
    tbl_w = table.find("w:tblPr/w:tblW", NS)
    tbl_ind = table.find("w:tblPr/w:tblInd", NS)
    widths = [int(x.get(f"{{{NS['w']}}}w")) for x in table.findall("w:tblGrid/w:gridCol", NS)]
    if tbl_w is None or int(tbl_w.get(f"{{{NS['w']}}}w")) != sum(widths):
        errors.append(f"Table {idx}: tblW/grid mismatch")
    if tbl_ind is None or tbl_ind.get(f"{{{NS['w']}}}w") != "120":
        errors.append(f"Table {idx}: tblInd mismatch")
    for ridx, row in enumerate(table.findall("w:tr", NS), 1):
        cells = row.findall("w:tc", NS)
        cell_widths = []
        for cell in cells:
            tcw = cell.find("w:tcPr/w:tcW", NS)
            cell_widths.append(int(tcw.get(f"{{{NS['w']}}}w")) if tcw is not None else -1)
        if cell_widths != widths:
            errors.append(f"Table {idx} row {ridx}: tcW/grid mismatch")
        height = row.find("w:trPr/w:trHeight", NS)
        if height is not None and height.get(f"{{{NS['w']}}}hRule") == "exact":
            errors.append(f"Table {idx} row {ridx}: fixed exact height")

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} abstracts={len(ABSTRACTS)}")
print("word_counts=" + ",".join(str(word_count(x["abstract"])) for x in ABSTRACTS))
if errors:
    print("AUDIT FAILED")
    for error in errors:
        print("-", error)
    sys.exit(1)
print("AUDIT PASSED")
